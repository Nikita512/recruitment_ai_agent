import pdfplumber
from docx import Document
import io

def extract_text_from_pdf(path_or_bytes):
    text = ''
    if isinstance(path_or_bytes, (bytes, bytearray)):
        with pdfplumber.open(io.BytesIO(path_or_bytes)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
    else:
        with pdfplumber.open(path_or_bytes) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ''
    return text

def extract_text_from_docx(path_or_bytes):
    text = ''
    if isinstance(path_or_bytes, (bytes, bytearray)):
        doc = Document(io.BytesIO(path_or_bytes))
    else:
        doc = Document(path_or_bytes)
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def extract_text_generic(fileobj, filename):
    content = fileobj.read()
    if filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.lower().endswith('.docx') or filename.lower().endswith('.doc'):
        return extract_text_from_docx(content)
    else:
        # fallback: try decode
        try:
            return content.decode('utf-8')
        except:
            return ''